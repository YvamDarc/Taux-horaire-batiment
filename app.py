import io
import datetime as dt

import streamlit as st
import pandas as pd
import altair as alt

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.pdfgen import canvas

st.set_page_config(
    page_title="Rapprochement CA / Achats / Heures – Bâtiment",
    layout="wide"
)

# =========================
# Utils
# =========================
def fmt_eur(x):
    try:
        return f"{float(x):,.0f} €".replace(",", " ")
    except Exception:
        return ""

def fmt_pct(x):
    try:
        return f"{float(x) * 100:.1f} %"
    except Exception:
        return ""

def safe_div(a, b):
    return a / b if b not in (0, None) else 0.0

def normalize_hours_df(df: pd.DataFrame) -> pd.DataFrame:
    """Normalise la table des heures (types + colonne heures facturables)."""
    if df is None or df.empty:
        df = pd.DataFrame([{"Personne": "", "Heures": 0.0, "Coef_production": 0.0}])

    out = df.copy()
    for col in ["Personne", "Heures", "Coef_production"]:
        if col not in out.columns:
            out[col] = "" if col == "Personne" else 0.0

    out["Personne"] = out["Personne"].astype(str).fillna("")
    out["Heures"] = pd.to_numeric(out["Heures"], errors="coerce").fillna(0.0)
    out["Coef_production"] = pd.to_numeric(out["Coef_production"], errors="coerce").fillna(0.0)

    out["Heures_facturables"] = out["Heures"] * out["Coef_production"]
    return out[["Personne", "Heures", "Coef_production", "Heures_facturables"]]

def compute_year(ca, achats, df_hours, taux_horaire, coef_refact):
    ca = float(ca or 0.0)
    achats = float(achats or 0.0)
    taux_horaire = float(taux_horaire or 0.0)
    coef_refact = float(coef_refact or 0.0)

    marge = ca - achats
    tx_marge = safe_div(marge, ca) if ca else 0.0

    dfh = normalize_hours_df(df_hours)
    heures = float(dfh["Heures"].sum())
    heures_fact = float(dfh["Heures_facturables"].sum())

    ca_theo_achats = achats * coef_refact
    ca_theo_heures = heures_fact * taux_horaire
    ca_theo_total = ca_theo_achats + ca_theo_heures

    ecart = ca - ca_theo_total
    ecart_pct = safe_div(ecart, ca) if ca else 0.0

    return {
        "marge": marge,
        "tx_marge": tx_marge,
        "dfh": dfh,
        "heures": heures,
        "heures_fact": heures_fact,
        "ca_theo_achats": ca_theo_achats,
        "ca_theo_heures": ca_theo_heures,
        "ca_theo_total": ca_theo_total,
        "ecart": ecart,
        "ecart_pct": ecart_pct,
    }

# =========================
# Excel (trame + import)
# =========================
REQUIRED_COLS = ["Personne", "Heures", "Coef_production"]

def make_template_excel_bytes() -> bytes:
    """Génère une trame Excel (2 onglets N et N-1) en mémoire."""
    df_n = pd.DataFrame(
        [
            {"Personne": "Ouvrier 1", "Heures": 140, "Coef_production": 0.75},
            {"Personne": "Ouvrier 2", "Heures": 152, "Coef_production": 0.70},
        ],
        columns=REQUIRED_COLS
    )
    df_n1 = pd.DataFrame(
        [
            {"Personne": "Ouvrier 1", "Heures": 138, "Coef_production": 0.72},
            {"Personne": "Ouvrier 2", "Heures": 150, "Coef_production": 0.68},
        ],
        columns=REQUIRED_COLS
    )

    buf = io.BytesIO()
    with pd.ExcelWriter(buf, engine="openpyxl") as writer:
        df_n.to_excel(writer, index=False, sheet_name="N")
        df_n1.to_excel(writer, index=False, sheet_name="N-1")
    return buf.getvalue()

def read_hours_sheet(xls: pd.ExcelFile, sheet_name: str):
    if sheet_name not in xls.sheet_names:
        return None
    df = pd.read_excel(xls, sheet_name=sheet_name)
    df.columns = [str(c).strip() for c in df.columns]
    missing = set(REQUIRED_COLS) - set(df.columns)
    if missing:
        raise ValueError(f"Onglet '{sheet_name}' : colonnes manquantes: {', '.join(sorted(missing))}")
    return df[REQUIRED_COLS].copy()

def load_hours_from_excel(uploaded_file):
    """Attend un xlsx avec onglet N (obligatoire) + onglet N-1 (optionnel)."""
    xls = pd.ExcelFile(uploaded_file)
    df_n = read_hours_sheet(xls, "N")
    if df_n is None:
        raise ValueError("L'onglet 'N' est obligatoire.")
    df_n1 = read_hours_sheet(xls, "N-1")
    return df_n, df_n1

# =========================
# Export Word / PDF
# =========================
def build_summary_payload(
    use_n1: bool,
    ca_n: float, achats_n: float, taux_horaire_n: float, coef_refact_n: float, res_n: dict,
    ca_n1: float | None, achats_n1: float | None, taux_horaire_n1: float | None, coef_refact_n1: float | None, res_n1: dict | None,
):
    """Structure stable pour exporter (Word/PDF)."""
    payload = {
        "date": dt.date.today().strftime("%d/%m/%Y"),
        "N": {
            "ca": ca_n,
            "achats": achats_n,
            "taux_horaire": taux_horaire_n,
            "coef_refact": coef_refact_n,
            "res": res_n,
        },
        "N-1": None,
        "use_n1": bool(use_n1 and res_n1 is not None),
    }
    if payload["use_n1"]:
        payload["N-1"] = {
            "ca": ca_n1,
            "achats": achats_n1,
            "taux_horaire": taux_horaire_n1,
            "coef_refact": coef_refact_n1,
            "res": res_n1,
        }
    return payload

def add_docx_kv_table(doc: Document, title: str, rows: list[tuple[str, str]]):
    doc.add_heading(title, level=2)
    table = doc.add_table(rows=1, cols=2)
    hdr = table.rows[0].cells
    hdr[0].text = "Indicateur"
    hdr[1].text = "Valeur"

    for k, v in rows:
        r = table.add_row().cells
        r[0].text = str(k)
        r[1].text = str(v)

def add_docx_hours_table(doc: Document, title: str, dfh: pd.DataFrame):
    doc.add_heading(title, level=2)
    df_show = dfh.copy()
    # arrondis propres
    for col in ["Heures", "Coef_production", "Heures_facturables"]:
        df_show[col] = pd.to_numeric(df_show[col], errors="coerce").fillna(0.0)

    table = doc.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text = "Personne"
    hdr[1].text = "Heures"
    hdr[2].text = "Coef production"
    hdr[3].text = "Heures facturables"

    for _, row in df_show.iterrows():
        r = table.add_row().cells
        r[0].text = str(row.get("Personne", ""))
        r[1].text = f"{float(row.get('Heures', 0.0)):.2f}"
        r[2].text = f"{float(row.get('Coef_production', 0.0)):.2f}"
        r[3].text = f"{float(row.get('Heures_facturables', 0.0)):.2f}"

def generate_docx_report(payload: dict) -> bytes:
    doc = Document()

    # Style simple
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_paragraph("Récapitulatif — Rapprochement CA / Achats / Heures (Bâtiment)")
    title.runs[0].bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Date : {payload['date']}").alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(" ")

    def section_for(label: str, block: dict):
        res = block["res"]
        doc.add_heading(f"Année {label}", level=1)

        add_docx_kv_table(
            doc,
            "Paramètres",
            [
                ("Taux horaire", f"{block['taux_horaire']:.2f} €/h"),
                ("Coef refacturation achats", f"{block['coef_refact']:.2f}"),
            ],
        )

        add_docx_kv_table(
            doc,
            "Données & résultats",
            [
                ("Chiffre d'affaires (réel)", fmt_eur(block["ca"])),
                ("Achats", fmt_eur(block["achats"])),
                ("Marge (CA - Achats)", fmt_eur(res["marge"])),
                ("Taux de marge", fmt_pct(res["tx_marge"])),
                ("Heures totales", f"{res['heures']:.2f} h"),
                ("Heures facturables", f"{res['heures_fact']:.2f} h"),
                ("CA théorique achats", fmt_eur(res["ca_theo_achats"])),
                ("CA théorique heures", fmt_eur(res["ca_theo_heures"])),
                ("CA théorique total", fmt_eur(res["ca_theo_total"])),
                ("Écart (réel - théorique)", fmt_eur(res["ecart"])),
                ("Écart (%)", fmt_pct(res["ecart_pct"])),
            ],
        )

        add_docx_hours_table(doc, "Détail heures par personne", res["dfh"])
        doc.add_page_break()

    section_for("N", payload["N"])
    if payload["use_n1"]:
        section_for("N-1", payload["N-1"])

    # Remove last page break if it exists (simple heuristic)
    # (optional; leave as-is if you don't care)
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()

def pdf_draw_kv(c: canvas.Canvas, x: float, y: float, key: str, val: str, key_w: float = 7.0*cm):
    c.setFont("Helvetica", 10)
    c.drawString(x, y, key)
    c.setFont("Helvetica-Bold", 10)
    c.drawString(x + key_w, y, val)

def pdf_draw_table_hours(c: canvas.Canvas, x: float, y: float, dfh: pd.DataFrame, max_rows: int = 28):
    # Header
    c.setFont("Helvetica-Bold", 10)
    c.drawString(x, y, "Personne")
    c.drawString(x + 8.5*cm, y, "Heures")
    c.drawString(x + 11.0*cm, y, "Coef")
    c.drawString(x + 13.0*cm, y, "Heures fact.")
    y -= 0.5*cm
    c.setLineWidth(0.5)
    c.line(x, y, x + 18.0*cm, y)
    y -= 0.35*cm

    c.setFont("Helvetica", 9)
    df_show = dfh.copy()
    for col in ["Heures", "Coef_production", "Heures_facturables"]:
        df_show[col] = pd.to_numeric(df_show[col], errors="coerce").fillna(0.0)

    rows = 0
    for _, r in df_show.iterrows():
        if rows >= max_rows:
            return y, True  # overflow
        c.drawString(x, y, str(r.get("Personne", ""))[:45])
        c.drawRightString(x + 10.3*cm, y, f"{float(r.get('Heures', 0.0)):.2f}")
        c.drawRightString(x + 12.6*cm, y, f"{float(r.get('Coef_production', 0.0)):.2f}")
        c.drawRightString(x + 17.8*cm, y, f"{float(r.get('Heures_facturables', 0.0)):.2f}")
        y -= 0.45*cm
        rows += 1

    return y, False

def generate_pdf_report(payload: dict) -> bytes:
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    width, height = A4

    def cover_page():
        c.setFont("Helvetica-Bold", 16)
        c.drawCentredString(width/2, height - 3.0*cm, "Récapitulatif — Rapprochement CA / Achats / Heures (Bâtiment)")
        c.setFont("Helvetica", 11)
        c.drawCentredString(width/2, height - 4.0*cm, f"Date : {payload['date']}")
        c.setFont("Helvetica", 10)
        c.drawString(2.0*cm, height - 6.0*cm,
                     "Ce document présente un comparatif entre CA réel et CA théorique (Achats refacturés + Heures facturées).")
        c.showPage()

    def year_page(label: str, block: dict):
        res = block["res"]
        x = 2.0*cm
        y = height - 2.3*cm

        c.setFont("Helvetica-Bold", 14)
        c.drawString(x, y, f"Année {label}")
        y -= 0.9*cm

        c.setFont("Helvetica-Bold", 12)
        c.drawString(x, y, "Paramètres")
        y -= 0.6*cm

        pdf_draw_kv(c, x, y, "Taux horaire", f"{block['taux_horaire']:.2f} €/h"); y -= 0.5*cm
        pdf_draw_kv(c, x, y, "Coef refacturation achats", f"{block['coef_refact']:.2f}"); y -= 0.8*cm

        c.setFont("Helvetica-Bold", 12)
        c.drawString(x, y, "Données & résultats")
        y -= 0.6*cm

        lines = [
            ("Chiffre d'affaires (réel)", fmt_eur(block["ca"])),
            ("Achats", fmt_eur(block["achats"])),
            ("Marge (CA - Achats)", fmt_eur(res["marge"])),
            ("Taux de marge", fmt_pct(res["tx_marge"])),
            ("Heures totales", f"{res['heures']:.2f} h"),
            ("Heures facturables", f"{res['heures_fact']:.2f} h"),
            ("CA théorique achats", fmt_eur(res["ca_theo_achats"])),
            ("CA théorique heures", fmt_eur(res["ca_theo_heures"])),
            ("CA théorique total", fmt_eur(res["ca_theo_total"])),
            ("Écart (réel - théorique)", fmt_eur(res["ecart"])),
            ("Écart (%)", fmt_pct(res["ecart_pct"])),
        ]
        for k, v in lines:
            pdf_draw_kv(c, x, y, k, v)
            y -= 0.5*cm

        y -= 0.2*cm
        c.setFont("Helvetica-Bold", 12)
        c.drawString(x, y, "Détail heures par personne")
        y -= 0.7*cm

        # Table hours, multi-pages if needed
        overflow = True
        dfh = res["dfh"]
        while overflow:
            y, overflow = pdf_draw_table_hours(c, x, y, dfh, max_rows=28)
            if overflow:
                c.showPage()
                y = height - 2.3*cm
                c.setFont("Helvetica-Bold", 12)
                c.drawString(x, y, f"Année {label} — Détail heures (suite)")
                y -= 0.9*cm

        c.showPage()

    cover_page()
    year_page("N", payload["N"])
    if payload["use_n1"]:
        year_page("N-1", payload["N-1"])

    c.save()
    return buf.getvalue()

# =========================
# Default session state
# =========================
if "hours_n" not in st.session_state:
    st.session_state["hours_n"] = pd.DataFrame(
        [
            {"Personne": "Ouvrier 1", "Heures": 140, "Coef_production": 0.75},
            {"Personne": "Ouvrier 2", "Heures": 140, "Coef_production": 0.70},
        ],
        columns=REQUIRED_COLS
    )

if "hours_n1" not in st.session_state:
    st.session_state["hours_n1"] = pd.DataFrame(
        [
            {"Personne": "Ouvrier 1", "Heures": 140, "Coef_production": 0.70},
            {"Personne": "Ouvrier 2", "Heures": 140, "Coef_production": 0.68},
        ],
        columns=REQUIRED_COLS
    )

# =========================
# UI
# =========================
st.title("Rapprochement CA / Achats / Heures — Bâtiment")
st.caption("Comparer le CA réel à un CA théorique (refact achats + facturation heures), en N et N-1.")

# =========================
# Sidebar – Paramètres
# =========================
with st.sidebar:
    st.header("Paramètres – Année N")
    taux_horaire_n = st.number_input("Taux horaire N (€/h)", min_value=0.0, value=55.0, step=1.0)
    coef_refact_n = st.number_input("Coef refact achats N", min_value=0.0, value=1.15, step=0.01, format="%.2f")

    st.divider()
    use_n1 = st.checkbox("Activer l’année N-1", value=False)

    if use_n1:
        st.header("Paramètres – Année N-1")
        taux_horaire_n1 = st.number_input("Taux horaire N-1 (€/h)", min_value=0.0, value=52.0, step=1.0)
        coef_refact_n1 = st.number_input("Coef refact achats N-1", min_value=0.0, value=1.12, step=0.01, format="%.2f")
    else:
        taux_horaire_n1 = None
        coef_refact_n1 = None

# =========================
# 0) Trame Excel + Import
# =========================
st.subheader("0️⃣ Trame Excel (téléchargement) & Import")

left, right = st.columns([1, 1])

with left:
    st.markdown("### Télécharger une trame Excel")
    template_bytes = make_template_excel_bytes()
    st.download_button(
        label="📄 Télécharger la trame (.xlsx)",
        data=template_bytes,
        file_name="trame_heures_batiment.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        use_container_width=True,
    )
    st.caption("Trame avec 2 onglets : **N** et **N-1** (colonnes : Personne, Heures, Coef_production).")

with right:
    st.markdown("### Importer une trame Excel")
    uploaded = st.file_uploader(
        "Importer un fichier .xlsx (onglet 'N' obligatoire, 'N-1' optionnel)",
        type=["xlsx"],
        accept_multiple_files=False
    )

    col_btn1, col_btn2 = st.columns([1, 1])
    with col_btn1:
        if uploaded is not None and st.button("🔄 Charger depuis l'Excel", use_container_width=True):
            try:
                df_n, df_n1 = load_hours_from_excel(uploaded)
                st.session_state["hours_n"] = df_n.copy()
                if df_n1 is not None:
                    st.session_state["hours_n1"] = df_n1.copy()
                st.success("Import OK ✅ Tableaux rechargés. Tu peux ensuite modifier manuellement.")
            except Exception as e:
                st.error(f"Import impossible : {e}")

    with col_btn2:
        with st.expander("Aperçu trame (exemple)"):
            st.write("**Onglet N** / **Onglet N-1** : mêmes colonnes")
            st.dataframe(
                pd.DataFrame(
                    [
                        {"Personne": "Ouvrier 1", "Heures": 140, "Coef_production": 0.75},
                        {"Personne": "Ouvrier 2", "Heures": 152, "Coef_production": 0.70},
                    ],
                    columns=REQUIRED_COLS
                ),
                use_container_width=True
            )

st.divider()

# =========================
# 1) CA / Achats
# =========================
st.subheader("1️⃣ Chiffre d’affaires et achats")

c1, c2 = st.columns(2)

with c1:
    st.markdown("### Année N")
    ca_n = st.number_input("CA N", min_value=0.0, value=300000.0, step=1000.0)
    achats_n = st.number_input("Achats N", min_value=0.0, value=150000.0, step=1000.0)

with c2:
    st.markdown("### Année N-1")
    if use_n1:
        ca_n1 = st.number_input("CA N-1", min_value=0.0, value=280000.0, step=1000.0)
        achats_n1 = st.number_input("Achats N-1", min_value=0.0, value=140000.0, step=1000.0)
    else:
        ca_n1 = None
        achats_n1 = None
        st.info("N-1 désactivé")

st.divider()

# =========================
# 2) Heures – Edition manuelle
# =========================
st.subheader("2️⃣ Heures par personne (modifiable manuellement)")

h1, h2 = st.columns(2)

with h1:
    st.markdown("### Ouvriers – N")
    st.session_state["hours_n"] = st.data_editor(
        st.session_state["hours_n"],
        num_rows="dynamic",
        use_container_width=True,
        key="editor_hours_n",
        column_config={
            "Personne": st.column_config.TextColumn(required=True),
            "Heures": st.column_config.NumberColumn(min_value=0.0, step=1.0),
            "Coef_production": st.column_config.NumberColumn(min_value=0.0, max_value=2.0, step=0.01, format="%.2f"),
        },
    )

with h2:
    st.markdown("### Ouvriers – N-1")
    if use_n1:
        if st.button("Copier le tableau N → N-1", use_container_width=True):
            st.session_state["hours_n1"] = st.session_state["hours_n"].copy()

        st.session_state["hours_n1"] = st.data_editor(
            st.session_state["hours_n1"],
            num_rows="dynamic",
            use_container_width=True,
            key="editor_hours_n1",
            column_config={
                "Personne": st.column_config.TextColumn(required=True),
                "Heures": st.column_config.NumberColumn(min_value=0.0, step=1.0),
                "Coef_production": st.column_config.NumberColumn(min_value=0.0, max_value=2.0, step=0.01, format="%.2f"),
            },
        )
    else:
        st.info("N-1 désactivé")

st.caption("Heures facturables = Heures × Coef de production.")
st.divider()

# =========================
# 3) Calculs
# =========================
res_n = compute_year(ca_n, achats_n, st.session_state["hours_n"], taux_horaire_n, coef_refact_n)
res_n1 = None
if use_n1 and ca_n1 is not None and achats_n1 is not None:
    res_n1 = compute_year(ca_n1, achats_n1, st.session_state["hours_n1"], taux_horaire_n1, coef_refact_n1)

# =========================
# 3) KPI
# =========================
st.subheader("3️⃣ Résultats")

k1, k2, k3, k4 = st.columns(4)
k1.metric("Marge N", fmt_eur(res_n["marge"]), fmt_pct(res_n["tx_marge"]))
k2.metric("Heures facturables N", f"{res_n['heures_fact']:.1f} h")
k3.metric("CA théorique N", fmt_eur(res_n["ca_theo_total"]))
k4.metric("Écart N (réel − théorique)", fmt_eur(res_n["ecart"]), fmt_pct(res_n["ecart_pct"]))

if res_n1 is not None:
    st.divider()
    k1b, k2b, k3b, k4b = st.columns(4)
    k1b.metric("Marge N-1", fmt_eur(res_n1["marge"]), fmt_pct(res_n1["tx_marge"]))
    k2b.metric("Heures facturables N-1", f"{res_n1['heures_fact']:.1f} h")
    k3b.metric("CA théorique N-1", fmt_eur(res_n1["ca_theo_total"]))
    k4b.metric("Écart N-1 (réel − théorique)", fmt_eur(res_n1["ecart"]), fmt_pct(res_n1["ecart_pct"]))

st.divider()

# =========================
# 4) Graphiques (Altair)
# =========================
st.subheader("4️⃣ Analyse graphique")

rows = [
    {"Année": "N", "Type": "CA réel", "Montant": float(ca_n)},
    {"Année": "N", "Type": "CA théorique", "Montant": float(res_n["ca_theo_total"])},
]
if res_n1 is not None:
    rows += [
        {"Année": "N-1", "Type": "CA réel", "Montant": float(ca_n1)},
        {"Année": "N-1", "Type": "CA théorique", "Montant": float(res_n1["ca_theo_total"])},
    ]
df_compare = pd.DataFrame(rows)

chart_ca = (
    alt.Chart(df_compare)
    .mark_bar()
    .encode(
        x=alt.X("Année:N", title="Année"),
        xOffset=alt.XOffset("Type:N"),
        y=alt.Y("Montant:Q", title="Montant (€)"),
        color=alt.Color("Type:N", legend=alt.Legend(title="")),
        tooltip=[alt.Tooltip("Année:N"), alt.Tooltip("Type:N"), alt.Tooltip("Montant:Q", format=",.0f")],
    )
)

labels_ca = (
    alt.Chart(df_compare)
    .mark_text(dy=-8)
    .encode(
        x=alt.X("Année:N"),
        xOffset=alt.XOffset("Type:N"),
        y=alt.Y("Montant:Q"),
        detail="Type:N",
        text=alt.Text("Montant:Q", format=",.0f"),
    )
)

gap_rows = [{"Année": "N", "Écart": float(res_n["ecart"])}]
if res_n1 is not None:
    gap_rows.append({"Année": "N-1", "Écart": float(res_n1["ecart"])})
df_gap = pd.DataFrame(gap_rows)

gap_bar = (
    alt.Chart(df_gap)
    .mark_bar()
    .encode(
        x=alt.X("Année:N", title="Année"),
        y=alt.Y("Écart:Q", title="Écart (€)"),
        color=alt.condition(
            alt.datum["Écart"] >= 0,
            alt.value("#2e7d32"),
            alt.value("#c62828"),
        ),
        tooltip=[alt.Tooltip("Année:N"), alt.Tooltip("Écart:Q", format=",.0f")],
    )
)

gap_zero = alt.Chart(pd.DataFrame({"y": [0]})).mark_rule().encode(y="y:Q")

gap_labels = (
    alt.Chart(df_gap)
    .mark_text(dy=-8)
    .encode(
        x="Année:N",
        y="Écart:Q",
        text=alt.Text("Écart:Q", format=",.0f"),
    )
)

comp_rows = [
    {"Année": "N", "Composant": "Achats / revente", "Montant": float(res_n["ca_theo_achats"])},
    {"Année": "N", "Composant": "Heures", "Montant": float(res_n["ca_theo_heures"])},
]
if res_n1 is not None:
    comp_rows += [
        {"Année": "N-1", "Composant": "Achats / revente", "Montant": float(res_n1["ca_theo_achats"])},
        {"Année": "N-1", "Composant": "Heures", "Montant": float(res_n1["ca_theo_heures"])},
    ]
df_comp = pd.DataFrame(comp_rows)

chart_comp = (
    alt.Chart(df_comp)
    .mark_bar()
    .encode(
        x=alt.X("Année:N", title="Année"),
        y=alt.Y("sum(Montant):Q", title="CA théorique (€)"),
        color=alt.Color("Composant:N", legend=alt.Legend(title="")),
        tooltip=[alt.Tooltip("Année:N"), alt.Tooltip("Composant:N"), alt.Tooltip("Montant:Q", format=",.0f")],
    )
)

labels_comp = (
    alt.Chart(df_comp)
    .mark_text(color="white")
    .encode(
        x="Année:N",
        y=alt.Y("Montant:Q", stack="zero"),
        detail="Composant:N",
        text=alt.Text("Montant:Q", format=",.0f"),
    )
)

colA, colB = st.columns(2)
with colA:
    st.markdown("### CA réel vs CA théorique")
    st.altair_chart((chart_ca + labels_ca).properties(height=340), use_container_width=True)

with colB:
    st.markdown("### Écart (réel − théorique)")
    st.altair_chart((gap_bar + gap_zero + gap_labels).properties(height=340), use_container_width=True)

st.markdown("### Composition du CA théorique")
st.altair_chart((chart_comp + labels_comp).properties(height=360), use_container_width=True)

# =========================
# 5) Export Word / PDF
# =========================
st.divider()
st.subheader("5️⃣ Export récap (Word / PDF)")

payload = build_summary_payload(
    use_n1=use_n1,
    ca_n=ca_n, achats_n=achats_n, taux_horaire_n=taux_horaire_n, coef_refact_n=coef_refact_n, res_n=res_n,
    ca_n1=ca_n1, achats_n1=achats_n1, taux_horaire_n1=taux_horaire_n1, coef_refact_n1=coef_refact_n1, res_n1=res_n1,
)

colW, colP = st.columns(2)

with colW:
    docx_bytes = generate_docx_report(payload)
    st.download_button(
        "📝 Télécharger le récap Word (.docx)",
        data=docx_bytes,
        file_name=f"recap_ca_achats_heures_{dt.date.today().strftime('%Y%m%d')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True,
    )

with colP:
    pdf_bytes = generate_pdf_report(payload)
    st.download_button(
        "🧾 Télécharger le récap PDF",
        data=pdf_bytes,
        file_name=f"recap_ca_achats_heures_{dt.date.today().strftime('%Y%m%d')}.pdf",
        mime="application/pdf",
        use_container_width=True,
    )

with st.expander("Détails des heures (avec heures facturables)"):
    st.markdown("#### N")
    st.dataframe(res_n["dfh"], use_container_width=True)
    if res_n1 is not None:
        st.markdown("#### N-1")
        st.dataframe(res_n1["dfh"], use_container_width=True)
